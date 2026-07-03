"""Pure KB parser + extraction gate + size guard for the Rehearsal trainer (Plan 04-01).

Livekit-free module mirroring ``agent/metrics.py`` / ``agent/persona.py``: frozen
module-level constants, ``@dataclass`` results, a ``_self_check()`` guarded by
``if __name__ == "__main__":`` (``python3 agent/kb/parse.py``). The whole
dispatch → normalize → extraction-gate → size-guard path is testable over fixture
bytes with NO livekit import — exactly like the metrics/persona pure cores.

Design rules baked in (CODE_PRINCIPLES + 04-PATTERNS Pattern A):
  * Boundary discipline (§4): ``parse`` is a named boundary — NO bare except, NO
    silent swallow. Every failure returns a typed ``KbParseError`` (the four
    ``reason`` values map 1:1 to a ``kb.state`` message); it NEVER raises across
    the boundary, so the agent's voice loop keeps running on a bad upload.
  * No magic values (§2): the gate/guard thresholds are named module constants.
  * The size guard measures EXTRACTED TOKENS, not file bytes — a scanned 200-page
    PDF is ~0 useful tokens, and a small text file can still distill-bust the
    frozen prefix budget. Tokens are the unit that couples to ``num_ctx``.
  * DOCX uses ``python-docx`` only — pymupdf4llm's Office support needs the paid
    PyMuPDF Pro (04-RESEARCH §3.2). PDF/TXT/MD are free via pymupdf4llm/stdlib.
  * No volatile data (no wall-clock reads, counters, or unique ids) so the module
    stays a pure, deterministic function of its inputs.

Heavy parser deps (``pymupdf4llm``, ``fitz``/``pymupdf``, ``python-docx``) are
imported LAZILY inside ``_extract`` so the pure dispatch/gate/guard path (and the
``.txt``/``.md`` stdlib path the self-check exercises) imports without them.
"""
from __future__ import annotations

import io
import sys
from dataclasses import dataclass

# Supported upload kinds (KB-01). Anything outside this set is an "unsupported"
# typed error — never a raise.
SUPPORTED: tuple[str, ...] = ("pdf", "txt", "md", "docx")

# Extraction-quality gate thresholds (Pattern A2 — detect empty/scanned/garbage
# BEFORE distillation). A scanned PDF extracts ~0 useful text or a low ratio of
# alphabetic characters.
ALPHA_RATIO_FLOOR: float = 0.50
MIN_USEFUL_CHARS: int = 32

# Size-guard thresholds, measured in ESTIMATED EXTRACTED TOKENS (Pattern A3).
# CHARS_PER_TOKEN is a cheap estimate (a real tokenizer is optional). These three
# are COUPLED to the distilled-brief budget + ``num_ctx`` pinned in 04-02/04-03:
# raising KB_MAX_TOKENS without bumping num_ctx would blow the frozen-prefix
# budget. Keep them in lockstep with BRIEF_TOKEN_BUDGET (04-02) and the Modelfile
# num_ctx (04-03).
CHARS_PER_TOKEN: int = 4
KB_WARN_TOKENS: int = 6000      # over → distill harder (KB-08 oversize_warn signal for 04-02)
KB_MAX_TOKENS: int = 24000      # over → reject this doc with a clear error (REL-03)

# Aggregate ceiling across ALL accepted docs in a session (M2). The per-doc
# KB_MAX_TOKENS guard bounds ONE upload, but ``ingest_kb`` re-distills the FULL
# concatenation of every accepted doc each time, so several docs (or many small
# ones) can push the distill INPUT past the effective Ollama context — Ollama then
# silently truncates the prompt (the GAP-1 truncation, multi-doc edition). Reject a
# new doc when the running session total would exceed this aggregate budget. Kept at
# KB_MAX_TOKENS so the multi-doc total can never exceed what a single max-size doc
# is already allowed to be (coupled to OLLAMA_CONTEXT_LENGTH the same way).
KB_AGGREGATE_MAX_TOKENS: int = KB_MAX_TOKENS

# RESOURCE ceilings enforced BEFORE the (memory-heavy) extraction step, so a large
# or maliciously-crafted upload cannot OOM-kill the worker before the token guard
# (which only measures EXTRACTED text) ever runs.
#   * KB_MAX_RAW_BYTES — cheap len(raw) ceiling on the on-the-wire upload, checked
#     before _extract buffers/parses it. Generous enough for real PDFs/DOCX.
#   * DOCX_MAX_UNCOMPRESSED_BYTES — sum of the .docx zip members' UNCOMPRESSED sizes,
#     checked before python-docx decompresses them. A tiny .docx whose document.xml
#     is a zip bomb expands to GBs in memory inside Document(...) — small on the wire
#     (so KB_MAX_RAW_BYTES does not catch it), huge once inflated. Reject by the zip
#     directory's declared sizes before handing the bytes to python-docx.
KB_MAX_RAW_BYTES: int = 25 * 1024 * 1024          # 25 MB upload ceiling
DOCX_MAX_UNCOMPRESSED_BYTES: int = 50 * 1024 * 1024  # 50 MB inflated-zip ceiling

# PDF page ceiling (F16), enforced on doc.page_count BEFORE pymupdf4llm.to_markdown
# walks every page. PDF content streams are Flate-compressed, so a small-on-the-wire
# PDF (within KB_MAX_RAW_BYTES) can declare ~100k pages; to_markdown runs per-page
# layout analysis via asyncio.to_thread while holding ingest_lock, and threads are
# UNCANCELLABLE — one pathological PDF pins a CPU indefinitely and blocks every later
# upload. The token guard only runs AFTER full extraction, so it can't help here.
# A real doc is far under this; over it → the oversize typed error (readable, too big).
PDF_MAX_PAGES: int = 2000

# Extension → kind map for the few cases where the name is authoritative.
_EXT_KIND: dict[str, str] = {
    "pdf": "pdf",
    "txt": "txt",
    "md": "md",
    "markdown": "md",
    "docx": "docx",
}

# MIME → kind map (fallback when the extension is missing/ambiguous).
_MIME_KIND: dict[str, str] = {
    "application/pdf": "pdf",
    "text/plain": "txt",
    "text/markdown": "md",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


class _OversizeExtraction(Exception):
    """Internal: a resource ceiling tripped DURING extraction (e.g. a DOCX zip bomb).

    Caught inside ``parse`` and converted to a typed ``oversize`` ``KbParseError`` —
    never raised across the public boundary.
    """


@dataclass
class ParsedDoc:
    """A successfully parsed document. ``text`` is normalized UTF-8 (BOM/smart-quote
    cleaned). ``oversize_warn`` is True when the token estimate is over
    ``KB_WARN_TOKENS`` but within ``KB_MAX_TOKENS`` — the KB-08 "distill harder"
    signal that 04-02's distiller reads.
    """

    name: str
    text: str
    token_estimate: int
    oversize_warn: bool


@dataclass
class KbParseError:
    """A typed parse failure. ``reason`` is one of the four matrix values; ``message``
    is the user-facing clear error surfaced via the ``kb.state`` attribute (REL-03).
    Returned, never raised across the parse boundary.
    """

    name: str
    reason: str        # "unsupported" | "scanned" | "corrupt" | "oversize"
    message: str


def parse(name: str, mime: str, raw: bytes) -> ParsedDoc | KbParseError:
    """Dispatch by type, normalize, run the extraction gate, then the size guard.

    Returns a ``ParsedDoc`` on success or a typed ``KbParseError`` for any of
    unsupported / scanned / corrupt / oversize. NEVER raises across this boundary
    (the parser-exception path is caught and converted to a ``corrupt`` error).
    """
    kind = _kind(name, mime)
    if kind not in SUPPORTED:
        return KbParseError(name, "unsupported", "Unsupported file type — use PDF/TXT/MD/DOCX")
    # Pre-parse byte ceiling (H1): reject an oversize upload BEFORE _extract buffers
    # and parses it into memory. The KB_MAX_TOKENS guard below only measures EXTRACTED
    # text, so without this a multi-hundred-MB upload could OOM the worker during
    # extraction before that guard ever runs.
    if len(raw) > KB_MAX_RAW_BYTES:
        return KbParseError(name, "oversize", "Too large for inline KB — trimmed/skipped")
    try:
        text = _extract(kind, raw)
    except _OversizeExtraction:
        # A resource ceiling tripped DURING extraction (e.g. a DOCX zip bomb whose
        # inflated size exceeds DOCX_MAX_UNCOMPRESSED_BYTES). Map to the oversize
        # reason, not corrupt — the file is readable, just too big to inline.
        return KbParseError(name, "oversize", "Too large for inline KB — trimmed/skipped")
    except Exception:
        # Named boundary (CODE_PRINCIPLES §4): any parser exception becomes a typed
        # error so the voice loop keeps running. No bare except, no silent swallow.
        return KbParseError(name, "corrupt", "Couldn't read this file")
    text = _normalize(text)
    gate = _extraction_gate(text)
    if gate is not None:
        gate.name = name
        return gate
    token_estimate = _estimate_tokens(text)
    if token_estimate > KB_MAX_TOKENS:
        return KbParseError(name, "oversize", "Too large for inline KB — trimmed/skipped")
    return ParsedDoc(
        name=name,
        text=text,
        token_estimate=token_estimate,
        oversize_warn=token_estimate > KB_WARN_TOKENS,
    )


def _kind(name: str, mime: str) -> str:
    """Resolve the document kind from the file extension, then the MIME type."""
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    if ext in _EXT_KIND:
        return _EXT_KIND[ext]
    return _MIME_KIND.get((mime or "").strip().lower(), "")


def _extract(kind: str, raw: bytes) -> str:
    """Extract raw text per kind. Heavy parser deps are imported lazily so the pure
    dispatch/gate/guard path imports without pymupdf4llm/python-docx.
    """
    if kind == "pdf":
        return _extract_pdf(raw)
    if kind == "docx":
        return _extract_docx(raw)
    # txt / md: decode bytes directly (errors replaced so a stray byte never raises).
    return raw.decode("utf-8", errors="replace")


def _pdf_pages_over_cap(page_count: int) -> bool:
    """PURE: True iff a PDF's declared page count exceeds PDF_MAX_PAGES (F16).

    Split out so the page-ceiling decision is unit-testable without PyMuPDF (which is
    a lazy, sandbox-absent dep) — mirrors the DOCX zip-directory size check shape.
    """
    return page_count > PDF_MAX_PAGES


def _extract_pdf(raw: bytes) -> str:
    """PDF → markdown via pymupdf4llm over an in-memory PyMuPDF document (no temp file).

    Page-count guard (F16): reject BEFORE to_markdown when doc.page_count exceeds
    PDF_MAX_PAGES. page_count is O(1) from the PDF page tree (no per-page layout
    work), so the pathological many-page PDF is bounced before the uncancellable
    per-page walk can pin a CPU under ingest_lock. Over-cap raises _OversizeExtraction
    (mapped to the oversize typed error by parse), exactly like the DOCX zip bomb.
    """
    import fitz  # PyMuPDF; lazy so the pure path imports without it
    import pymupdf4llm

    doc = fitz.open(stream=raw, filetype="pdf")
    try:
        if _pdf_pages_over_cap(doc.page_count):
            raise _OversizeExtraction(doc.page_count)
        return pymupdf4llm.to_markdown(doc)
    finally:
        # PyMuPDF holds an open handle on the in-memory stream; close it so no
        # resource lingers (KB-06 ephemerality — nothing persists).
        doc.close()


def _extract_docx(raw: bytes) -> str:
    """DOCX → text via python-docx (NOT pymupdf4llm — Office needs paid PyMuPDF Pro).

    Joins paragraph text AND explicitly iterates table cells: tables extract out
    of order if ignored (Pitfall 14), so they are walked separately and appended.
    """
    import zipfile  # stdlib; used to inspect the .docx zip directory before inflating

    # Zip-bomb guard (H2): a .docx is a ZIP and python-docx decompresses its members
    # with no ratio/size limit, so a small-on-the-wire .docx whose document.xml is a
    # zip bomb inflates to GBs inside Document(...) and OOM-kills the worker (the H1
    # raw-byte cap does NOT catch it — the upload is tiny). Sum the zip directory's
    # DECLARED uncompressed sizes and bail BEFORE the heavy python-docx import + parse.
    # Runs on stdlib `zipfile` only, ahead of the lazy `from docx import Document`.
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            uncompressed = sum(zi.file_size for zi in zf.infolist())
    except zipfile.BadZipFile:
        uncompressed = 0  # not a valid zip → let python-docx raise the real error below
    if uncompressed > DOCX_MAX_UNCOMPRESSED_BYTES:
        raise _OversizeExtraction(uncompressed)

    from docx import Document  # python-docx; lazy import

    doc = Document(io.BytesIO(raw))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _normalize(text: str) -> str:
    """Strip the BOM and normalize common smart-quote/dash artifacts to ASCII.

    Cheap, deterministic, no volatile data. Ensures a clean UTF-8 ``str`` for the
    gate and the (later) byte-stable distilled brief.
    """
    cleaned = text.lstrip("\ufeff")
    for src, dst in _SMART_QUOTE_MAP.items():
        cleaned = cleaned.replace(src, dst)
    return cleaned


# Common Unicode punctuation → ASCII (BOM handled separately in _normalize).
_SMART_QUOTE_MAP: dict[str, str] = {
    "\u2018": "'", "\u2019": "'",      # single curly quotes
    "\u201c": '"', "\u201d": '"',      # double curly quotes
    "\u2013": "-", "\u2014": "-",      # en/em dash
    "\u2026": "...",                    # ellipsis
}


def _extraction_gate(text: str) -> KbParseError | None:
    """Detect empty/scanned/garbage extraction (Pattern A2). Returns a ``scanned``
    error (with an empty ``name`` for the caller to fill) or ``None`` when the text
    looks like real prose.
    """
    stripped = text.strip()
    if len(stripped) < MIN_USEFUL_CHARS:
        return KbParseError("", "scanned", "Couldn't extract text (looks scanned)")
    alpha_ratio = sum(c.isalpha() for c in stripped) / len(stripped)
    if alpha_ratio < ALPHA_RATIO_FLOOR:
        return KbParseError("", "scanned", "Couldn't extract text (looks scanned)")
    return None


def _estimate_tokens(text: str) -> int:
    """Cheap token estimate over EXTRACTED text (not file bytes): chars // CHARS_PER_TOKEN."""
    return len(text) // CHARS_PER_TOKEN


def _self_check() -> None:
    """Pure-stdlib check over fixture bytes (``python3 agent/kb/parse.py``).

    Exercises the full dispatch + extraction gate + size guard + oversize_warn path
    using only the txt/md stdlib branch (no pymupdf4llm/python-docx needed), mirror-
    ing ``metrics.py``/``persona.py``. Asserts the four typed-error reasons and the
    oversize_warn signal. Not part of the runtime path.
    """
    prose = b"The quick brown fox jumps over the lazy dog. " * 4
    ok = parse("notes.txt", "text/plain", prose)
    assert isinstance(ok, ParsedDoc), f"expected ParsedDoc, got {ok!r}"
    assert ok.name == "notes.txt" and ok.token_estimate > 0
    assert ok.oversize_warn is False, "small prose should not warn"

    unsupported = parse("archive.bin", "application/octet-stream", b"\x7fELF")
    assert isinstance(unsupported, KbParseError) and unsupported.reason == "unsupported"

    scanned = parse("scan.txt", "text/plain", b"\x00\x01\x02")
    assert isinstance(scanned, KbParseError) and scanned.reason == "scanned"
    assert scanned.name == "scan.txt", "gate error must carry the file name"

    oversize_raw = ("word " * (KB_MAX_TOKENS * CHARS_PER_TOKEN)).encode("utf-8")
    oversize = parse("huge.txt", "text/plain", oversize_raw)
    assert isinstance(oversize, KbParseError) and oversize.reason == "oversize"

    # Pre-parse byte cap (H1): a raw upload over KB_MAX_RAW_BYTES is rejected as
    # oversize BEFORE extraction (cheap len(raw) check, no parse attempted).
    too_many_bytes = b"a" * (KB_MAX_RAW_BYTES + 1)
    byte_capped = parse("big.txt", "text/plain", too_many_bytes)
    assert isinstance(byte_capped, KbParseError) and byte_capped.reason == "oversize", (
        "raw byte cap must reject before extraction"
    )

    # Mid-size text: tokens between WARN and MAX → ParsedDoc with oversize_warn True.
    warn_chars = (KB_WARN_TOKENS + (KB_MAX_TOKENS - KB_WARN_TOKENS) // 2) * CHARS_PER_TOKEN
    warn_raw = ("a" * warn_chars).encode("utf-8")
    warn = parse("mid.txt", "text/plain", warn_raw)
    assert isinstance(warn, ParsedDoc) and warn.oversize_warn is True, "mid-size should warn"

    print("kb.parse _self_check OK", file=sys.stderr)


if __name__ == "__main__":
    _self_check()
